#!/usr/bin/env python3
"""由本目录的 Markdown 生成公文体 .docx（依赖 python-docx）。

用法：python3 build_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

BASE = Path(__file__).resolve().parent

SOURCES = [
    ("01-tenglong-window-details-and-opinion.md",
     "01-tenglong-window-details-and-opinion.docx"),
    ("02-telecom-lishui-window-details-and-opinion.md",
     "02-telecom-lishui-window-details-and-opinion.docx"),
]

BODY_FONT = "仿宋"
HEI = "黑体"
KAI = "楷体"
TITLE_FONT = "方正小标宋简体"
ASCII_FONT = "Times New Roman"


def set_run_font(run, east_asian, size_pt, bold=False):
    run.font.name = ASCII_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), east_asian)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(28)
    pf.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def add_page_number_footer(doc):
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        set_run_font(run, BODY_FONT, 10.5)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._element.append(begin)
        run._element.append(instr)
        run._element.append(end)


def add_paragraph(doc, text, *, style=None, font=BODY_FONT, size=12, bold=False,
                  align=None, indent_chars=0, space_before=0, space_after=0):
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_chars:
        pf.first_line_indent = Pt(size * indent_chars)
    # 标签（如【问题】）加粗，其余正文常规
    match = re.match(r"^(【[^】]+】)(.*)$", text, flags=re.S)
    if match and not bold:
        label = para.add_run(match.group(1))
        set_run_font(label, HEI, size, bold=False)
        rest = para.add_run(match.group(2))
        set_run_font(rest, font, size, bold=False)
    else:
        run = para.add_run(text)
        set_run_font(run, font, size, bold=bold)
    return para


def clean_inline(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def build(md_path, docx_path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    add_page_number_footer(doc)

    subtitle_done = False
    for raw in lines:
        line = clean_inline(raw)
        if not line or line.startswith("---"):
            continue

        if line.startswith("#### "):
            add_paragraph(doc, line[5:].strip(), font=BODY_FONT, size=12,
                          bold=True, indent_chars=2, space_before=6)
        elif line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), font=KAI, size=14,
                          bold=True, indent_chars=2, space_before=8)
        elif line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), font=HEI, size=15,
                          indent_chars=2, space_before=12, space_after=4)
        elif line.startswith("# "):
            add_paragraph(doc, line[2:].strip(), font=TITLE_FONT, size=20,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        elif not subtitle_done and line.startswith("（") and line.endswith("）"):
            add_paragraph(doc, line, font=KAI, size=12,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            subtitle_done = True
        else:
            add_paragraph(doc, line, indent_chars=2)

    doc.save(docx_path)
    return docx_path


def main():
    for md_name, docx_name in SOURCES:
        out = build(BASE / md_name, BASE / docx_name)
        print(f"generated: {out.name}")


if __name__ == "__main__":
    main()
