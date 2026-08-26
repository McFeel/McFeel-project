#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_merged_book.py —— 将产品报告 docx 深拷进佛山书稿，并把书稿版面套成产品报告实际样式。

用法：
    python3 build_merged_book.py 源书稿.docx --report 产品报告.docx --output 融合后书稿.docx
    python3 build_merged_book.py 源书稿.docx --dry-run        # 仅打印将替换的块（无需 --report）

行为：
    1. 只读打开源书稿与产品报告两个 docx；
    2. 在书稿正文中定位最后一处同时含「第三篇附」与「产品介绍」的段落作为薄占位块
       起点（跳过目录样式条目与含「对接：」的交叉引用行）；块终点为下一个同级或
       更高级标题、或「第×篇／附录×」；
    3. 把产品报告正文的全部块级内容（段落 w:p、表格 w:tbl，含 a:blip／v:imagedata
       图片与外部超链接关系）深拷到插入点，替换薄占位块；报告中的结构化文档域
       （w:sdt，如自动目录域）不拷入；
    4. 版面照产品报告实际样式：其 section 设置（pgSz／pgMar／cols／docGrid）套到
       书稿全部 section；其 styles.xml（docDefaults、latentStyles 及全部样式定义）
       合并进书稿，同 styleId 以报告为准；报告 numbering 中的列表定义重映射 id
       后并入书稿（不与书稿既有编号冲突）；
    5. 写入 --output 指定的新文件。

安全保证：
    * 两个输入文件均只读，绝不修改或覆盖；运行前后校验二者 SHA-256 一致；
    * --output 与任一输入路径相同时拒绝执行；
    * --output 已存在时拒绝覆盖，除非显式传入 --force；
    * 仅依赖 python-docx 与 Python 标准库。
"""

import argparse
import copy
import hashlib
import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.numbering import NumberingPart

# ------------------------------------------------------- 占位块定位特征

BLOCK_END_RE = re.compile(r"^(附录\s*[A-DＡ-Ｄ]|第[一二三四五六七八九十百零]+[篇部])")
TOC_STYLE_RE = re.compile(r"^(TOC|Contents|目录)\s*\d*$", re.IGNORECASE)
XREF_MARKS = ("对接：", "对接:")          # 交叉引用行，不参与定位
BLOCK_START_KEYS = ("第三篇附", "产品介绍")  # 须同时出现

# sectPr 子元素 schema 顺序（用于按位插入）
SECT_SEQ = (
    "w:headerReference", "w:footerReference", "w:footnotePr", "w:endnotePr",
    "w:type", "w:pgSz", "w:pgMar", "w:paperSrc", "w:pgBorders", "w:lnNumType",
    "w:pgNumType", "w:cols", "w:formProt", "w:vAlign", "w:noEndnote",
    "w:titlePg", "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
    "w:printerSettings",
)
SECT_COPY_TAGS = ("w:pgSz", "w:pgMar", "w:cols", "w:docGrid")

# VML 命名空间（python-docx 的 nsmap 未收录 v: 前缀，此处用完整 URI）
VML_IMAGEDATA = "{http://schemas.openxmlformats.org/vml}imagedata"


def log(msg):
    print(f"【build_merged_book】{msg}")


def sha256_of(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def heading_level(paragraph):
    """返回段落的标题级别（Title 视为 0）；非标题段落返回 None。"""
    name = paragraph.style.name or ""
    if name == "Title":
        return 0
    m = re.match(r"^Heading (\d+)$", name)
    if m:
        return int(m.group(1))
    return None


def is_toc_paragraph(p):
    """目录条目段落（Word 自动目录样式 TOC 1—9／目录 1—9）不参与占位块定位。"""
    return bool(TOC_STYLE_RE.match((p.style.name or "").strip()))


def find_annex_block(doc):
    """定位薄占位块，返回 (start_idx, end_idx)，end_idx 为块后第一段（可为 len）。

    起点：正文中最后一处同时含「第三篇附」与「产品介绍」的段落；
    跳过目录样式条目与含「对接：」的交叉引用行。
    """
    paras = doc.paragraphs
    candidates = []
    for i, p in enumerate(paras):
        if is_toc_paragraph(p):
            continue
        text = p.text.strip()
        if not text or any(mark in text for mark in XREF_MARKS):
            continue
        if all(key in text for key in BLOCK_START_KEYS):
            candidates.append(i)
    if not candidates:
        return None
    start = candidates[-1]
    start_level = heading_level(paras[start])
    end = len(paras)
    for j in range(start + 1, len(paras)):
        p = paras[j]
        text = p.text.strip()
        level = heading_level(p)
        if level is not None and (start_level is None or level <= start_level):
            end = j
            break
        if BLOCK_END_RE.match(text):
            end = j
            break
    return start, end


# ------------------------------------------------------- numbering 合并

def get_numbering_element(doc):
    """返回文档 numbering 部件的 w:numbering 元素；无 numbering 部件则返回 None。"""
    for rel in doc.part.rels.values():
        if rel.reltype == RT.NUMBERING and not rel.is_external:
            return rel.target_part.element
    return None


def ensure_numbering_element(doc):
    """取得书稿 numbering 元素；书稿无 numbering 部件时新建一个。"""
    el = get_numbering_element(doc)
    if el is not None:
        return el
    el = parse_xml("<w:numbering %s/>" % nsdecls("w"))
    part = NumberingPart(
        PackURI("/word/numbering.xml"), CT.WML_NUMBERING, el, doc.part.package
    )
    doc.part.relate_to(part, RT.NUMBERING)
    return el


def merge_numbering(book_doc, report_doc):
    """把报告 numbering 的列表定义并入书稿（id 重映射），返回 {报告numId: 书稿numId}。"""
    r_el = get_numbering_element(report_doc)
    if r_el is None:
        return {}
    b_el = ensure_numbering_element(book_doc)
    max_num = max([int(n.get(qn("w:numId"))) for n in b_el.findall(qn("w:num"))] + [0])
    max_abs = max(
        [int(n.get(qn("w:abstractNumId"))) for n in b_el.findall(qn("w:abstractNum"))] + [-1]
    )
    abs_map = {}
    first_num = b_el.find(qn("w:num"))
    for absn in r_el.findall(qn("w:abstractNum")):
        max_abs += 1
        abs_map[absn.get(qn("w:abstractNumId"))] = str(max_abs)
        new_absn = copy.deepcopy(absn)
        new_absn.set(qn("w:abstractNumId"), str(max_abs))
        if first_num is not None:
            first_num.addprevious(new_absn)   # abstractNum 须排在 num 之前
        else:
            b_el.append(new_absn)
    num_map = {}
    for num in r_el.findall(qn("w:num")):
        max_num += 1
        num_map[num.get(qn("w:numId"))] = str(max_num)
        new_num = copy.deepcopy(num)
        new_num.set(qn("w:numId"), str(max_num))
        ref = new_num.find(qn("w:abstractNumId"))
        if ref is not None and ref.get(qn("w:val")) in abs_map:
            ref.set(qn("w:val"), abs_map[ref.get(qn("w:val"))])
        b_el.append(new_num)
    return num_map


def remap_numids(element, num_map):
    """把元素子树中 w:numId 的引用按映射表改写。"""
    if not num_map:
        return
    for num_id in element.iter(qn("w:numId")):
        old = num_id.get(qn("w:val"))
        if old in num_map:
            num_id.set(qn("w:val"), num_map[old])


# ------------------------------------------------------- 样式与版面套用

def merge_styles(book_doc, report_doc, num_map):
    """报告 styles.xml 合并进书稿：docDefaults／latentStyles／全部样式定义，
    同 styleId 以报告为准；样式中的 numId 引用一并重映射。返回合并的样式数。"""
    b_styles = book_doc.styles.element
    r_styles = report_doc.styles.element
    for tag in ("w:docDefaults", "w:latentStyles"):
        r_el = r_styles.find(qn(tag))
        if r_el is None:
            continue
        new_el = copy.deepcopy(r_el)
        b_el = b_styles.find(qn(tag))
        if b_el is not None:
            b_styles.replace(b_el, new_el)
        else:
            b_styles.insert(0, new_el)
    book_by_id = {s.get(qn("w:styleId")): s for s in b_styles.findall(qn("w:style"))}
    count = 0
    for r_s in r_styles.findall(qn("w:style")):
        new_s = copy.deepcopy(r_s)
        remap_numids(new_s, num_map)
        old = book_by_id.get(new_s.get(qn("w:styleId")))
        if old is not None:
            b_styles.replace(old, new_s)
        else:
            b_styles.append(new_s)
        count += 1
    return count


def apply_section_props(book_doc, report_doc):
    """报告首个 section 的 pgSz／pgMar／cols／docGrid 套到书稿全部 section。"""
    r_sect_pr = report_doc.sections[0]._sectPr
    count = 0
    for b_sec in book_doc.sections:
        b_sect_pr = b_sec._sectPr
        for tag in SECT_COPY_TAGS:
            r_el = r_sect_pr.find(qn(tag))
            if r_el is None:
                continue
            new_el = copy.deepcopy(r_el)
            b_el = b_sect_pr.find(qn(tag))
            if b_el is not None:
                b_sect_pr.replace(b_el, new_el)
            else:
                idx = SECT_SEQ.index(tag)
                b_sect_pr.insert_element_before(new_el, *SECT_SEQ[idx + 1:])
        count += 1
    return count


# ------------------------------------------------------- 内容深拷

def extract_report_content(report_doc):
    """报告正文的全部块级元素（段落与表格）；w:sdt（自动目录域等）不拷入。"""
    body = report_doc.element.body
    return [el for el in body if el.tag in (qn("w:p"), qn("w:tbl"))]


def import_image(book_doc, report_doc, r_id, counters):
    """把报告中的图片关系导入书稿包（按内容 SHA1 去重），返回书稿侧新 rId。"""
    rel = report_doc.part.rels.get(r_id)
    if rel is None or rel.is_external:
        log(f"警告：图片关系 {r_id} 缺失或为外部链接，保留原样")
        return None
    blob = rel.target_part.blob
    new_r_id, _image = book_doc.part.get_or_add_image(io.BytesIO(blob))
    counters["img"] += 1
    return new_r_id


def fix_relationships(book_doc, report_doc, element, counters):
    """改写深拷元素中的关系引用：图片（a:blip／v:imagedata）与外部超链接。"""
    for blip in element.iter(qn("a:blip")):
        for attr in (qn("r:embed"), qn("r:link")):
            r_id = blip.get(attr)
            if not r_id:
                continue
            new_r_id = import_image(book_doc, report_doc, r_id, counters)
            if new_r_id:
                blip.set(attr, new_r_id)
    for imagedata in element.iter(VML_IMAGEDATA):
        r_id = imagedata.get(qn("r:id"))
        if not r_id:
            continue
        new_r_id = import_image(book_doc, report_doc, r_id, counters)
        if new_r_id:
            imagedata.set(qn("r:id"), new_r_id)
    for hyperlink in element.iter(qn("w:hyperlink")):
        r_id = hyperlink.get(qn("r:id"))
        if not r_id:
            continue
        rel = report_doc.part.rels.get(r_id)
        if rel is None:
            continue
        if rel.is_external:
            new_r_id = book_doc.part.relate_to(rel.target_ref, rel.reltype, is_external=True)
            hyperlink.set(qn("r:id"), new_r_id)
        else:
            log(f"警告：跳过指向报告包内部的超链接关系 {r_id}")


def prepare_element(book_doc, report_doc, src_el, num_map, counters):
    new_el = copy.deepcopy(src_el)
    if src_el.tag == qn("w:p"):
        counters["p"] += 1
    elif src_el.tag == qn("w:tbl"):
        counters["tbl"] += 1
    remap_numids(new_el, num_map)
    fix_relationships(book_doc, report_doc, new_el, counters)
    return new_el


def replace_block(book_doc, start, end, new_elements):
    """删除书稿中 [start, end) 段的薄占位块（含其中表格），插入深拷内容。"""
    paras = book_doc.paragraphs
    body = book_doc.element.body
    start_el = paras[start]._element
    end_el = paras[end]._element if end < len(paras) else None
    removed = 0
    el = start_el
    while el is not None and el is not end_el and el.tag != qn("w:sectPr"):
        nxt = el.getnext()
        body.remove(el)
        removed += 1
        el = nxt
    sect_pr = body.find(qn("w:sectPr"))
    for new_el in new_elements:
        if end_el is not None:
            end_el.addprevious(new_el)
        elif sect_pr is not None:
            sect_pr.addprevious(new_el)
        else:
            body.append(new_el)
    return removed


# ------------------------------------------------------- 主流程

def main(argv=None):
    parser = argparse.ArgumentParser(
        description="将产品报告 docx 深拷进佛山书稿并套用其版面（绝不覆盖输入文件）"
    )
    parser.add_argument("source", help="源书稿 docx 路径（只读，绝不修改）")
    parser.add_argument("--report", help="产品报告 docx 路径（只读；--dry-run 时不需要）")
    parser.add_argument("--output", help="融合后新文件路径（必需，除非 --dry-run）")
    parser.add_argument("--dry-run", action="store_true", help="仅打印将替换的块，不写文件")
    parser.add_argument("--force", action="store_true", help="允许覆盖已存在的输出文件")
    args = parser.parse_args(argv)

    src = Path(args.source).resolve()
    if not src.is_file():
        log(f"错误：源书稿不存在：{src}")
        return 2

    report_path = None
    if args.report:
        report_path = Path(args.report).resolve()
        if not report_path.is_file():
            log(f"错误：产品报告不存在：{report_path}")
            return 2

    out = None
    if not args.dry_run:
        if report_path is None:
            log("错误：缺少 --report（或使用 --dry-run 仅定位占位块）")
            return 2
        if not args.output:
            log("错误：缺少 --output（或使用 --dry-run）")
            return 2
        out = Path(args.output).resolve()
        if out == src or out == report_path:
            log("错误：输出路径与输入相同，已拒绝执行（绝不覆盖原件）")
            return 2
        if out.exists() and not args.force:
            log(f"错误：输出文件已存在：{out}（如需覆盖请显式加 --force）")
            return 2

    src_hash_before = sha256_of(src)
    report_hash_before = sha256_of(report_path) if report_path else None
    book = Document(str(src))

    found = find_annex_block(book)
    if found is None:
        log("错误：未找到薄占位块。定位规则：正文最后一处同时含「第三篇附」与"
            "「产品介绍」的段落（跳过目录样式条目与含「对接：」的交叉引用行）")
        return 3
    start, end = found
    paras = book.paragraphs
    log(f"定位到占位块：第 {start + 1} 段至第 {end} 段（共 {end - start} 段）")
    log(f"块首：「{paras[start].text.strip()[:60]}」")
    if end < len(paras):
        log(f"块后首段（保留）：「{paras[end].text.strip()[:60]}」")
    else:
        log("块后无内容（占位块位于文档末尾）")

    if args.dry_run:
        log("dry-run，将替换的块内容如下：")
        for p in paras[start:end]:
            print(f"    {p.text}")
        log("dry-run 结束，未写出任何文件")
        return 0

    report = Document(str(report_path))

    num_map = merge_numbering(book, report)
    log(f"numbering 合并完成：重映射 {len(num_map)} 个列表编号")
    n_styles = merge_styles(book, report, num_map)
    log(f"样式合并完成：报告 {n_styles} 个样式定义已套入书稿（同 styleId 以报告为准）")
    n_sections = apply_section_props(book, report)
    log(f"页面设置完成：报告 pgSz／pgMar／cols／docGrid 已套到书稿 {n_sections} 个 section")

    src_elements = extract_report_content(report)
    counters = {"p": 0, "tbl": 0, "img": 0}
    new_elements = [
        prepare_element(book, report, el, num_map, counters) for el in src_elements
    ]
    log(
        f"报告内容深拷完成：段落 {counters['p']} 个、表格 {counters['tbl']} 个、"
        f"图片关系 {counters['img']} 处"
    )

    removed = replace_block(book, start, end, new_elements)
    log(f"替换完成：删除原薄块 {removed} 个块级元素，插入报告内容 {len(new_elements)} 个")

    out.parent.mkdir(parents=True, exist_ok=True)
    book.save(str(out))

    if sha256_of(src) != src_hash_before:
        log("严重错误：源书稿校验值发生变化，已中止")
        return 4
    if report_hash_before and sha256_of(report_path) != report_hash_before:
        log("严重错误：产品报告校验值发生变化，已中止")
        return 4
    log("两个输入文件校验均未变（SHA-256 一致）")
    log(f"已写出新文件：{out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
