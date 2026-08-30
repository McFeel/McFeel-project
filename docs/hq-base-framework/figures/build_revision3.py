#!/usr/bin/env python3
"""Replace F1–F12 placeholders in revision 2 and save revision 3."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
FIGURES = Path(__file__).resolve().parent
DELIVERABLES = ROOT / "docs" / "hq-base-framework" / "deliverables"
SOURCE = DELIVERABLES / "南网总部基地建设框架方案_2026-08-30_修订2.docx"
OUTPUT = DELIVERABLES / "南网总部基地建设框架方案_2026-08-30_修订3.docx"


FIGURE_SPECS = {
    "F1": ("F01-政策束与三层战略收拢.png", "图 1-1 政策束与三层战略收拢图", 15.2),
    "F4": ("F04-既有园区现状与约束.png", "图 2-1 园区现状与约束一页图", 15.2),
    "F2": ("F02-一核六边总体架构.png", "图 3-1 一核六边总图", 15.0),
    "F3": ("F03-战略资产四件套与复制路径.png", "图 3-2 战略资产四件套与“内部验证—行业复制”路径图", 15.2),
    "F5": ("F05-目标三档与边界决策.png", "图 4-1 目标三档台阶与边界决策图", 15.2),
    "F6": ("F06-六边技术族行动矩阵.png", "图 5-1 六边×技术族矩阵总表", 15.2),
    "F9": ("F09-柔性资源与实时电价路线.png", "图 5-2 柔性资源—实时电价对接技术路线图", 15.2),
    "F7": ("F07-一个平台两类功能智慧架构.png", "图 5-3 “一个平台、两类功能”智慧架构图", 15.2),
    "F10": ("F10-普惠人人参与人人受益.png", "图 5-4 普惠“人人参与、人人受益”三环图", 15.2),
    "F8": ("F08-信息港做法到总部基地落点.png", "图 6-1 信息港做法→总部基地落点映射对照表", 15.2),
    "F11": ("F11-实施路线与三件先行.png", "图 7-1 实施路线图", 15.2),
    "F12": ("F12-证据等级与待核清单.png", "图 A-1 证据等级与待核清单表", 15.2),
}


REFERENCES = {
    "F1": ("四层政策束最终仍收拢到原有三层战略主线", "政策束与三层战略收拢关系见图 1-1。"),
    "F4": ("图示占位：F4", "五项硬约束及其方案含义见图 2-1。"),
    "F2": ("基地作为明牌，提供第二增长曲线所需的证据与信用", "一核六边总体架构见图 3-1，战略资产四件套与复制路径见图 3-2。"),
    "F5": ("中国建科院宣讲材料所载雄安案例", "目标档位与核算边界的前后关系见图 4-1。"),
    "F6": ("本章结论是：六边不是六组口号", "六边技术族、首批动作及资产沉淀关系见图 5-1。"),
    "F9": ("首批动作：冷站与水系统诊断进场", "柔性资源与实时电价的对接路线见图 5-2。"),
    "F7": ("首批动作：完成平台总体设计及数据主权条款", "“一个平台、两类功能”的完整架构见图 5-3。"),
    "F10": ("首批动作：设计个人—团队—园区碳账规则", "普惠三环及其共同底线见图 5-4。"),
    "F8": ("据交流纪要，相关平台支持本地部署但需按客户设备重训", "七项信息港做法、总部基地落点及差异化见图 6-1。"),
    "F11": ("第一阶段服务 9 月中旬方向性汇报", "三件先行、决策节点与实施节奏见图 7-1。"),
    "F12": ("本附录结论是：证据等级用于约束材料如何进入结论和指标", "证据使用规则与全量待核事项见图 A-1。"),
}


def set_run_font(run, size=9, color="44545E", bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def append_reference(document: Document, marker: str, sentence: str):
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            if sentence not in paragraph.text:
                run = paragraph.add_run(sentence)
                set_run_font(run, size=11, color="000000")
            return
    raise RuntimeError(f"Reference anchor not found: {marker}")


def replace_placeholder(document: Document, code: str, filename: str, caption: str, width_cm: float):
    marker = f"图示占位：{code}"
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(marker):
            for run in list(paragraph.runs):
                paragraph._p.remove(run._r)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt = paragraph.paragraph_format
            fmt.left_indent = Cm(0)
            fmt.right_indent = Cm(0)
            fmt.first_line_indent = Cm(0)
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1
            fmt.keep_together = True
            picture_run = paragraph.add_run()
            picture_run.add_picture(str(FIGURES / filename), width=Cm(width_cm))
            picture_run.add_break(WD_BREAK.LINE)
            caption_run = paragraph.add_run(caption)
            set_run_font(caption_run, size=9, color="344955", bold=True)
            return
    raise RuntimeError(f"Placeholder not found: {code}")


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    for filename, _, _ in FIGURE_SPECS.values():
        if not (FIGURES / filename).exists():
            raise FileNotFoundError(FIGURES / filename)

    document = Document(SOURCE)
    for _, (marker, sentence) in REFERENCES.items():
        append_reference(document, marker, sentence)
    for code, (filename, caption, width_cm) in FIGURE_SPECS.items():
        replace_placeholder(document, code, filename, caption, width_cm)

    document.core_properties.title = "南网总部基地建设框架方案（修订3）"
    document.core_properties.subject = "F1–F12 正文插图完成版"
    document.save(OUTPUT)
    print(f"Saved {OUTPUT}")


if __name__ == "__main__":
    main()
