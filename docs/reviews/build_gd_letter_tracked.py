#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_gd_letter_tracked.py

读入批注稿结构（gd_letter_annotated_structure.json：原稿段落 + 已审 20 条批注对应的
keep/del/ins 分段），生成带修订痕迹（w:ins / w:del，Track Changes）的 docx。

要点：
- 修订作者固定为「文秘」（w:author），时间取结构文件 meta.revision_date。
- 删除用 w:del + w:delText，插入用 w:ins + w:t；整段插入时同时把段落标记
  （pilcrow）标为插入（w:pPr/w:rPr/w:ins），接受修订后段落保留。
- 字体段落匹配原稿：标题黑体、一级标题黑体、二级标题楷体、正文/主送/落款仿宋
  （西文 Times New Roman），插入文字沿用所在段落原格式（即相邻原格式）。
- 只读输入文件，只写新输出文件，绝不覆盖任何输入文件。

用法：
    python3 build_gd_letter_tracked.py            # 生成 docx
    python3 build_gd_letter_tracked.py --verify   # 生成并校验修订痕迹与文字口径
"""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_STRUCTURE = SCRIPT_DIR / "gd_letter_annotated_structure.json"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# 接受修订后正文中不得出现的表述（应只存在于 w:del 删除痕迹中）
FORBIDDEN_IN_ACCEPTED = [
    "由南网能源公司承接",
    "无缝续签",
    "无缝接续",
    "恳请",
    "首先",
    "严格落实",
    "深度参与",
    "减少多供应商协调管理成本",
    "无法续签",
    "设备运维",
    "特此汇报",
    "工作汇报",
    "作专题汇报",
    "基层单位",
    "广东电网各单位",
    "2027-2028",  # 年份须用一字线 2027—2028年度
    "改造+运维",
]

# 接受修订后正文中必须出现的表述
REQUIRED_IN_ACCEPTED = [
    "关于商请将相关业务纳入2027—2028年度办公节能框架采购的函",
    "广东电网有限责任公司：",
    "南方电网综合能源股份有限公司（以下简称南网能源）",
    "BOO（建设—拥有—运营）",
    "厨房运维、空调运维",
    "2027—2028年度",
    "商请",
    "专此函达，请予研酌。",
    "依法依规参与竞争",
    "中选单位",
    "贵司所属单位",
    "（口径待核）",
    "（建设内容待核，到期日待核）",
    "（规模待核）",
    "项目清单及到期日待核",
    "缺少框架内采购渠道",
    "框架期内启动相关项目采购",
    "资产处置按原合同约定另行办理",
    "已签约或中选项目",
    "按合同及贵司技术标准",
    "    年  月  日",
]


def _set_fonts(rpr: "OxmlElement", style: dict) -> None:
    """按段落原格式设置中西文字体、字号、加粗。"""
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), style["ascii"])
    rfonts.set(qn("w:hAnsi"), style["ascii"])
    rfonts.set(qn("w:eastAsia"), style["eastAsia"])
    rfonts.set(qn("w:cs"), style["ascii"])
    rpr.append(rfonts)
    if style.get("bold"):
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))
    half_pt = str(int(round(style["size_pt"] * 2)))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), half_pt)
    rpr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), half_pt)
    rpr.append(szcs)


def _make_run(text: str, style: dict, deleted: bool) -> "OxmlElement":
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    _set_fonts(rpr, style)
    run.append(rpr)
    node = OxmlElement("w:delText" if deleted else "w:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    return run


class RevisionIds:
    def __init__(self) -> None:
        self._next = 1

    def __call__(self) -> str:
        rid = str(self._next)
        self._next += 1
        return rid


def _wrap(kind: str, run: "OxmlElement", rid: str, author: str, date: str) -> "OxmlElement":
    """把 run 包进 w:ins 或 w:del。"""
    el = OxmlElement(f"w:{kind}")
    el.set(qn("w:id"), rid)
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    el.append(run)
    return el


def _apply_paragraph_format(p, style: dict, inserted_mark: bool, rid: str, author: str, date: str) -> None:
    # 子元素顺序须符合 CT_PPr schema：spacing → ind → jc → rPr，否则 Word 可能报无法读取
    ppr = p._p.get_or_add_pPr()

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(round(style.get("line_pt", 28) * 20))))
    spacing.set(qn("w:lineRule"), "exact")
    ppr.append(spacing)

    indent_chars = int(style.get("first_line_indent_chars", 0))
    if indent_chars:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLineChars"), str(indent_chars * 100))
        ind.set(qn("w:firstLine"), str(int(round(style["size_pt"] * indent_chars * 20))))
        ppr.append(ind)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), {"left": "left", "center": "center", "right": "right", "justify": "both"}[style["align"]])
    ppr.append(jc)

    if inserted_mark:
        # 整段插入：段落标记也标为插入，接受修订后该段保留
        pmark_rpr = OxmlElement("w:rPr")
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), rid)
        ins.set(qn("w:author"), author)
        ins.set(qn("w:date"), date)
        pmark_rpr.append(ins)
        ppr.append(pmark_rpr)


def _setup_page(doc: Document) -> None:
    """A4，公文页边距（上 3.7cm / 下 3.5cm / 左 2.8cm / 右 2.6cm）。"""
    section = doc.sections[0]
    section.page_width = 11906 * 635  # twips -> EMU
    section.page_height = 16838 * 635
    section.top_margin = int(3.7 * 566.93) * 635
    section.bottom_margin = int(3.5 * 566.93) * 635
    section.left_margin = int(2.8 * 566.93) * 635
    section.right_margin = int(2.6 * 566.93) * 635


def _enable_track_changes(doc: Document) -> None:
    """开启修订跟踪。CT_Settings 中 w:trackChanges 须位于 w:defaultTabStop 之前。"""
    settings = doc.settings.element
    if settings.find(qn("w:trackChanges")) is not None:
        return
    track = OxmlElement("w:trackChanges")
    anchor = settings.find(qn("w:defaultTabStop"))
    if anchor is not None:
        anchor.addprevious(track)
    else:
        settings.insert(0, track)


def check_integrity(structure: dict) -> None:
    """keep+del 分段拼接必须等于原稿原文，确保未静默改动原文。"""
    for para in structure["paragraphs"]:
        reconstructed = "".join(s["text"] for s in para["segments"] if s["t"] in ("keep", "del"))
        if reconstructed != para["original"]:
            raise ValueError(
                f"段落 {para['id']} 分段拼接与原稿不一致：\n  拼接: {reconstructed!r}\n  原稿: {para['original']!r}"
            )
        if para.get("inserted_paragraph") and any(s["t"] != "ins" for s in para["segments"]):
            raise ValueError(f"段落 {para['id']} 标记为整段插入，但含非 ins 分段")


def build_docx(structure: dict, output_path: Path) -> None:
    meta = structure["meta"]
    author = meta.get("author", "文秘")
    date = meta.get("revision_date", "2026-08-27T00:00:00Z")
    styles = structure["styles"]
    rid = RevisionIds()

    doc = Document()
    _setup_page(doc)
    _enable_track_changes(doc)

    # 去掉模板自带空段落
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)

    for para in structure["paragraphs"]:
        style = styles[para["style"]]
        p = doc.add_paragraph()
        inserted_paragraph = bool(para.get("inserted_paragraph"))
        _apply_paragraph_format(p, style, inserted_paragraph, rid(), author, date)

        for seg in para["segments"]:
            kind, text = seg["t"], seg["text"]
            if not text:
                continue
            if kind == "keep":
                p._p.append(_make_run(text, style, deleted=False))
            elif kind == "ins":
                p._p.append(_wrap("ins", _make_run(text, style, deleted=False), rid(), author, date))
            elif kind == "del":
                p._p.append(_wrap("del", _make_run(text, style, deleted=True), rid(), author, date))
            else:
                raise ValueError(f"未知分段类型: {kind!r}")

    doc.save(str(output_path))


def verify_docx(output_path: Path) -> list[str]:
    """校验 docx：修订痕迹、作者、字体、文字口径。返回检查报告行。"""
    from lxml import etree

    report: list[str] = []
    with zipfile.ZipFile(output_path) as zf:
        xml = zf.read("word/document.xml")
    root = etree.fromstring(xml)

    ins_nodes = root.findall(f".//{{{W_NS}}}ins")
    del_nodes = root.findall(f".//{{{W_NS}}}del")
    report.append(f"w:ins 节点数: {len(ins_nodes)}")
    report.append(f"w:del 节点数: {len(del_nodes)}")
    assert ins_nodes, "未发现任何 w:ins 插入痕迹"
    assert del_nodes, "未发现任何 w:del 删除痕迹"

    authors = {n.get(qn("w:author")) for n in ins_nodes + del_nodes}
    report.append(f"修订作者集合: {sorted(authors)}")
    assert authors == {"文秘"}, f"修订作者必须为「文秘」，实际: {authors}"

    # 接受修订后的正文 = 所有 w:t（w:delText 是独立标签，不计入）
    accepted = "".join(t.text or "" for t in root.findall(f".//{{{W_NS}}}t"))
    deleted = "".join(t.text or "" for t in root.findall(f".//{{{W_NS}}}delText"))
    report.append(f"接受修订后正文字数: {len(accepted)}；删除痕迹字数: {len(deleted)}")

    for phrase in FORBIDDEN_IN_ACCEPTED:
        assert phrase not in accepted, f"禁用表述仍存在于接受后正文: {phrase}"
    report.append(f"禁用表述检查通过（{len(FORBIDDEN_IN_ACCEPTED)} 项均未出现在接受后正文）")

    for phrase in REQUIRED_IN_ACCEPTED:
        assert phrase in accepted, f"必需表述缺失: {phrase}"
    report.append(f"必需表述检查通过（{len(REQUIRED_IN_ACCEPTED)} 项均在接受后正文）")

    for phrase in ("由南网能源公司承接", "无缝续签", "恳请", "首先", "特此汇报，恳请贵司予以支持。"):
        assert phrase in deleted, f"应被删除的表述未出现在删除痕迹中: {phrase}"
    report.append("关键删除项均保留在 w:del 痕迹中")

    fonts = {rf.get(qn("w:eastAsia")) for rf in root.findall(f".//{{{W_NS}}}rFonts")}
    report.append(f"中文字体使用: {sorted(f for f in fonts if f)}")
    assert "黑体" in fonts and "仿宋" in fonts, f"字体不匹配原稿格式: {fonts}"

    # python-docx 回读，确认文件结构有效
    Document(str(output_path))
    report.append("python-docx 回读成功，docx 结构有效")
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="生成带修订痕迹（w:ins/w:del）的修订稿 docx")
    parser.add_argument("--structure", type=Path, default=DEFAULT_STRUCTURE, help="批注稿结构 JSON（只读输入）")
    parser.add_argument("--output", type=Path, default=None, help="输出 docx 路径（默认取结构文件 meta.output_filename）")
    parser.add_argument("--verify", action="store_true", help="生成后校验修订痕迹与文字口径")
    args = parser.parse_args()

    structure_path = args.structure.resolve()
    with open(structure_path, encoding="utf-8") as f:  # 只读，绝不写输入文件
        structure = json.load(f)

    output_path = args.output or (structure_path.parent / structure["meta"]["output_filename"])
    output_path = output_path.resolve()
    if output_path == structure_path:
        raise SystemExit("输出路径与输入文件相同，拒绝覆盖输入文件")

    check_integrity(structure)
    build_docx(structure, output_path)
    print(f"已生成: {output_path}")

    if args.verify:
        for line in verify_docx(output_path):
            print(f"[verify] {line}")
        print("[verify] 全部检查通过")
    return 0


if __name__ == "__main__":
    sys.exit(main())
