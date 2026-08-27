#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""夹具测试：不依赖原稿，用 python-docx 造一份含全部 20 个锚点的样例正文，
跑 add_gd_report_comments，验证 comments 部件与钉句标记是否正确生成。

用法：python3 docs/reviews/test_add_gd_report_comments.py
（无需第三方测试框架，断言失败即非零退出。）
"""

from __future__ import annotations

import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

import add_gd_report_comments as mod
from add_gd_report_comments import COMMENT_SPECS, FIRST_PARAGRAPH, add_comments

W = mod.W

# 每条问题对应的一段样例正文，句中含该条锚点（覆盖“或/斜杠”给出的候选之一）。
FIXTURE_PARAGRAPHS = [
    # 开篇段：既覆盖第5条“开篇无主送”（钉首段），也含第20条“首先衷心感谢”备选之一。
    "首先衷心感谢广东电网长期以来对我司办公节能业务的支持与信任。",
    "为进一步深化合作，现就2027-2028年度办公节能框架采购有关情况报告如下。",
    "该批业务拟由南网能源公司承接，并推进相关落地工作。",
    "我司将严格按照贵司制度流程参与框架采购，确保合规。",
    "现有多项合同将于2028年集中到期，若不提前谋划将影响服务连续性。",
    "部分合同到期后按现行安排无法续签，存在服务中断风险。",
    "统一采购有利于减少多供应商协调管理成本，提升整体效率。",
    "我司将严格落实各项管理要求，加强过程管控。",
    "在实施中着力保障改造质量、运维响应时效，让用户满意。",
    "拟对粤电大厦及清城供电局等物业开展节能改造。",
    "重点包括大楼供冷系统的节能升级与优化。",
    "同步覆盖基层单位的用能管理与节能诊断。",
    "并延伸至设备运维环节，形成闭环管理。",
    "（一）办公节能各细分业务",
    "拟构建2027-2028办公节能框架采购体系，统筹推进。",
    "本次拟首次采用BOO模式开展合作。",
    "总体看，相关业务增量拓展空间较大，前景可期。",
    "为保障服务不中断，恳请上级支持做到无缝续签。",
    "特此汇报。",
]


def build_fixture(path: Path) -> None:
    doc = Document()
    for text in FIXTURE_PARAGRAPHS:
        doc.add_paragraph(text)
    doc.save(str(path))


def _q(tag: str) -> str:
    return f"{{{W}}}{tag}"


def main() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        src = tmpdir / "fixture_src.docx"
        dst = tmpdir / "fixture_commented.docx"
        build_fixture(src)

        specs = add_comments(src, dst)

        # 1) 所有 20 条都命中锚点。
        assert len(specs) == 20, f"应有 20 条，实际 {len(specs)}"
        for spec in specs:
            assert spec.matched_anchor, f"第{spec.cid}条未命中锚点"

        with zipfile.ZipFile(dst) as zf:
            names = set(zf.namelist())
            comments_xml = zf.read("word/comments.xml")
            document_xml = zf.read("word/document.xml")
            ct_xml = zf.read("[Content_Types].xml")
            rels_xml = zf.read("word/_rels/document.xml.rels")

        # 2) comments 部件存在且含 20 条批注，author=文秘，正文写了严重程度+改法。
        assert "word/comments.xml" in names, "缺少 comments.xml 部件"
        croot = etree.fromstring(comments_xml)
        comments = croot.findall(_q("comment"))
        assert len(comments) == 20, f"comments.xml 应有 20 条，实际 {len(comments)}"
        seen_ids = set()
        for c in comments:
            cid = int(c.get(f"{{{W}}}id"))
            author = c.get(f"{{{W}}}author")
            seen_ids.add(cid)
            assert author == "文秘", f"author 应为 文秘，实际 {author}"
            body = "".join(t.text or "" for t in c.iter(_q("t")))
            assert body.startswith("【"), f"批注 {cid} 未以严重程度开头：{body[:12]}"
            assert "改法：" in body, f"批注 {cid} 缺改法：{body[:20]}"
        assert seen_ids == set(range(1, 21)), f"批注 id 应为 1..20，实际 {sorted(seen_ids)}"

        # 3) document.xml 里有 20 组 commentRangeStart/End 与 20 个 commentReference。
        droot = etree.fromstring(document_xml)
        starts = droot.findall(f".//{_q('commentRangeStart')}")
        ends = droot.findall(f".//{_q('commentRangeEnd')}")
        refs = droot.findall(f".//{_q('commentReference')}")
        assert len(starts) == 20, f"commentRangeStart 应 20，实际 {len(starts)}"
        assert len(ends) == 20, f"commentRangeEnd 应 20，实际 {len(ends)}"
        assert len(refs) == 20, f"commentReference 应 20，实际 {len(refs)}"
        start_ids = sorted(int(e.get(f"{{{W}}}id")) for e in starts)
        assert start_ids == list(range(1, 21)), f"range id 应 1..20，实际 {start_ids}"

        # 4) 正文一个字都没改：拼接全部 w:t 文字应与原夹具逐字一致。
        def all_text(xml: bytes) -> str:
            r = etree.fromstring(xml)
            body = r.find(f"{{{W}}}body")
            out = []
            for p in body.findall(f"{{{W}}}p"):
                out.append("".join(t.text or "" for t in p.iter(_q("t"))))
            return "\n".join(out)

        with zipfile.ZipFile(src) as zf:
            orig_text = all_text(zf.read("word/document.xml"))
        # 批注稿的正文段落里，comments.xml 不参与；只看 document.xml body。
        new_text = all_text(document_xml)
        assert orig_text == new_text, "正文文字被改动了！\n原：%r\n新：%r" % (
            orig_text,
            new_text,
        )

        # 5) 内容类型与关系都已登记。
        assert b"comments+xml" in ct_xml, "[Content_Types].xml 未登记 comments 类型"
        assert b"comments" in rels_xml, "document.xml.rels 未登记 comments 关系"

        # 6) python-docx 能重新打开产物（结构未损坏）。
        Document(str(dst))

        # 7) 第5条钉在首段（开篇无主送）。
        spec5 = next(s for s in specs if s.cid == 5)
        assert spec5.matched_anchor == FIRST_PARAGRAPH, "第5条应钉首段"

    print("OK：20 条批注全部命中，comments 部件、钉句标记、正文零改动均通过。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
