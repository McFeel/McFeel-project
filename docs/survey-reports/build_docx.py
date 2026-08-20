#!/usr/bin/env python3
"""把调研报告的 Markdown 正文转成公文体例的 Word 文档。

体例依据党政机关公文格式的通行做法：A4 页面，上37毫米、下35毫米、左28毫米、
右26毫米页边距；标题黑体居中；正文仿宋_GB2312 三号，首行缩进2字，行距固定值29磅；
层次依次为一、（一）1.，分别用黑体、楷体_GB2312、仿宋_GB2312加粗；页码居中。

用法：python3 docs/survey-reports/build_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

DOC_DIR = Path(__file__).resolve().parent
SOURCE = DOC_DIR / "2026-08-南网总部基地绿色近零碳智慧园区项目调研报告.md"
TARGET = DOC_DIR / "2026-08-南网总部基地绿色近零碳智慧园区项目调研报告.docx"

FANGSONG = "仿宋_GB2312"
HEITI = "黑体"
KAITI = "楷体_GB2312"

SIZE_TITLE = Pt(18)  # 小二号，取此字号可使本报告标题在一行内排完
SIZE_BODY = Pt(16)  # 三号
SIZE_TABLE = Pt(10.5)  # 五号
# 四列宽度合计156毫米，即A4去掉左右页边距后的可用宽度
TABLE_COL_WIDTHS = (Mm(22), Mm(40), Mm(47), Mm(47))
LINE_BODY = Pt(29)  # 每页22行的固定行距

COVER_END = "<!-- 封面结束 -->"
META_RE = re.compile(r"^\*\*(.+?)\*\*：(.*)$")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def set_font(run, name: str, size: Pt, bold: bool = False) -> None:
    """同时设置西文与中文字体，否则 Word 里中文会回落到默认字体。"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def add_paragraph(
    doc,
    text: str,
    *,
    font: str = FANGSONG,
    size: Pt = SIZE_BODY,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: bool = True,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
    line_spacing: Pt | None = LINE_BODY,
):
    """写入一段，按 **粗体** 切分 run，保留全角标点原样。"""
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.alignment = align
    fmt.space_before = space_before
    fmt.space_after = space_after
    if line_spacing is not None:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = line_spacing
    if first_line_indent:
        fmt.first_line_indent = size * 2

    pos = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > pos:
            set_font(para.add_run(text[pos : match.start()]), font, size, bold)
        set_font(para.add_run(match.group(1)), font, size, True)
        pos = match.end()
    if pos < len(text):
        set_font(para.add_run(text[pos:]), font, size, bold)
    if not text:
        set_font(para.add_run(""), font, size, bold)
    return para


def set_default_style(doc) -> None:
    """把 Normal 样式设为仿宋_GB2312 三号，避免未显式设字体处回落。"""
    style = doc.styles["Normal"]
    style.font.name = FANGSONG
    style.font.size = SIZE_BODY
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FANGSONG)


def configure_page(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)


def restart_page_numbering(section) -> None:
    """正文另起一节并从第1页重新编号，使封面不占页码。

    w:pgNumType 在 w:sectPr 中的位置受 schema 约束，必须排在 w:cols 之前，
    直接 append 会被 Word 忽略。
    """
    sect_pr = section._sectPr
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), "1")
    anchor = sect_pr.find(qn("w:cols"))
    if anchor is not None:
        anchor.addprevious(pg_num)
    else:
        sect_pr.append(pg_num)


def add_page_number_footer(section) -> None:
    """页码居中，采用公文常见的 — 1 — 形式。

    须先断开与上一节的链接，否则写入的是封面那一节的页脚。
    """
    section.footer.is_linked_to_previous = False
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def run_with(text: str):
        run = para.add_run(text)
        set_font(run, "宋体", Pt(14))
        return run

    run_with("— ")
    field = para.add_run()
    set_font(field, "宋体", Pt(14))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field._element.append(begin)
    field._element.append(instr)
    field._element.append(end)
    run_with(" —")


def build_cover(doc, title: str, subtitle: str, meta: list[tuple[str, str]]) -> None:
    for _ in range(6):
        add_paragraph(doc, "", first_line_indent=False, line_spacing=None)
    add_paragraph(
        doc,
        title,
        font=HEITI,
        size=SIZE_TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        line_spacing=None,
        space_after=Pt(18),
    )
    if subtitle:
        add_paragraph(
            doc,
            subtitle,
            font=KAITI,
            size=SIZE_BODY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            line_spacing=None,
        )
    for _ in range(8):
        add_paragraph(doc, "", first_line_indent=False, line_spacing=None)
    for label, value in meta:
        add_paragraph(
            doc,
            f"{label}：{value}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            line_spacing=None,
        )
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section)
    restart_page_numbering(body_section)
    add_page_number_footer(body_section)


def repeat_header_row(row) -> None:
    """跨页时重复表头。"""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def build_table(doc, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    repeat_header_row(table.rows[0])
    widths = TABLE_COL_WIDTHS[: len(rows[0])]
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.cell(r, c)
            cell.width = widths[c]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_font(
                para.add_run(cell_text),
                HEITI if r == 0 else FANGSONG,
                SIZE_TABLE,
                bold=False,
            )
    add_paragraph(doc, "", first_line_indent=False, space_before=Pt(6))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= {"-", ":", " "} for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def main() -> int:
    if not SOURCE.exists():
        print(f"未找到源文件：{SOURCE}")
        return 1

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_page(doc.sections[0])
    set_default_style(doc)

    cover_title = ""
    cover_subtitle = ""
    cover_meta: list[tuple[str, str]] = []
    in_cover = True
    i = 0
    trailing_meta: list[tuple[str, str]] = []
    after_last_rule = False

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line == COVER_END:
            build_cover(doc, cover_title, cover_subtitle, cover_meta)
            in_cover = False
            i += 1
            continue

        if not line:
            i += 1
            continue

        if in_cover:
            if line.startswith("# "):
                cover_title = line[2:].strip()
            elif line.startswith("## "):
                cover_subtitle = line[3:].strip()
            else:
                m = META_RE.match(line)
                if m:
                    cover_meta.append((m.group(1), m.group(2)))
            i += 1
            continue

        if set(line) <= {"-"} and len(line) >= 3:
            after_last_rule = True
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                build_table(doc, rows)
            continue

        if line.startswith("#### "):
            add_paragraph(
                doc,
                line[5:].strip(),
                font=FANGSONG,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(6),
            )
        elif line.startswith("### "):
            add_paragraph(
                doc,
                line[4:].strip(),
                font=KAITI,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(8),
            )
        elif line.startswith("## "):
            add_paragraph(
                doc,
                line[3:].strip(),
                font=HEITI,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(14),
                space_after=Pt(4),
            )
        else:
            m = META_RE.match(line)
            if after_last_rule and m:
                trailing_meta.append((m.group(1), m.group(2)))
            else:
                add_paragraph(doc, line)
        i += 1

    if trailing_meta:
        add_paragraph(doc, "", first_line_indent=False, space_before=Pt(12))
        for label, value in trailing_meta:
            align = (
                WD_ALIGN_PARAGRAPH.JUSTIFY if label == "附" else WD_ALIGN_PARAGRAPH.RIGHT
            )
            add_paragraph(
                doc,
                f"{label}：{value}" if label != "附" else f"附：{value}",
                align=align,
                first_line_indent=(label == "附"),
            )

    doc.save(TARGET)
    print(f"已生成 {TARGET.name}")
    print(f"段落数 {len(doc.paragraphs)}，表格数 {len(doc.tables)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
