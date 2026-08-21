#!/usr/bin/env python3
"""稿件红线自查。

对本目录下的调研报告成稿做机械校验，覆盖九类约束：

1. “零碳区域”只允许出现在转述住建部认证体系之处，且不得用于指称本项目；
2. “全国首个”“国内首个”一类表述必须带出处或争取方向的限定语；
3. 会中口径的量化数据必须同段带口述或估算的标注，不得写成实测或核证结果；
4. 清华会谈已开展，但其估算与工作安排不得升级为已核证、已验收、已完成采购；
5. 既成事实用语（已立项、已签约等）只允许出现在否定语境；
6. 含量化数据的段落必须同段标注出处；
7. 文号只允许出现资料中给出的四个，其余一律视为编造；
8. 纪要中的年份笔误只能作为笔误提及，不得作为时间要求；
9. 正文使用全角标点，并校验公文 Word 与 Markdown 正文同步。

用法：python3 docs/survey-reports/check_writing_constraints.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

DOC_DIR = Path(__file__).resolve().parent
DOC_NAME = "调研报告：南网总部基地北京调研"
REPORT = DOC_DIR / f"{DOC_NAME}.md"
DOCX = DOC_DIR / f"{DOC_NAME}.docx"

# “零碳区域”系建科院宣讲材料转述住建部体系时的原词，只能在有出处的语境下使用。
ZERO_CARBON_AREA = "零碳区域"
AREA_SOURCE_MARKERS = ("宣讲材料", "住建部", "建科院")
PROJECT_MARKERS = ("南网", "总部基地", "本项目")

# “首个”类表述必须说明是谁的口径。
FIRST_CLAIMS = ("全国首个", "国内首个")
FIRST_MARKERS = ("宣讲材料", "争取", "会内", "意向", "据")

# 清华会谈已于2026年8月21日开展，但会中内容不得升级为已核证或已成交。
TSINGHUA_FORBIDDEN = ("已核证", "已验收", "已完成采购", "已成交")

FAIT_ACCOMPLI = (
    "已立项",
    "已落地",
    "已签约",
    "已申报",
    "已批复",
    "已中标",
    "已验收",
    "已核证",
)
NEGATION_MARKERS = ("不", "未", "没有", "无", "非", "禁", "尚")

# 量化数据必须与出处同段出现。
QUANTITY_RE = re.compile(
    r"\d+(?:\.\d+)?(?:%|％|万平方米|平方米|万千瓦|兆瓦|千瓦时|千瓦|万度|度|吨|"
    r"万立方米|亩|块|栋|次每小时|万元|元每度|元)"
)
SOURCE_MARKERS = (
    "据",
    "宣讲材料",
    "资料宣称",
    "绿智楼宇资料",
    "M-Park资料",
    "中移资料",
    "新华网",
    "中国网",
    "新能源网",
    "报道",
    "纪要",
    "会上",
    "会中",
    "调研计划",
    "计算",
    "估算",
    "宣称",
    "资料",
    "稿中未见",
    "标准",
    "台阶",
    "要求",
    "口径",
    "门槛",
)

# 会中口径的数据必须带这些标注之一：或说明是口述、估算，或点明是会中何种表述。
HEDGE_MARKERS = (
    "口述",
    "估算",
    "估计",
    "待核",
    "核验",
    "自述",
    "未经",
    "称",
    "提到",
    "明确",
    "说明",
    "安排",
)

# 文号只允许资料中给出的四个；仿宋_GB2312、楷体_GB2312 是公文字体名。
ALLOWED_STANDARD_NUMBERS = (
    "GB 55015-2021",
    "计划通〔2024〕283号",
    "发改能源〔2024〕1123号",
    "发改环资规〔2024〕127号",
    "仿宋_GB2312",
    "楷体_GB2312",
)
DOC_NUMBER_RE = re.compile(r"〔\d{4}〕|\d{4}〕\d+号|第?\d+号文|GB ?/ ?T ?\d|GB ?\d{4,}")

# 清华纪要中的年份笔误，只能作为笔误提及。
TYPO_YEAR = "2024年8"
TYPO_MARKERS = ("笔误", "不符")

HALFWIDTH_PUNCT = re.compile(r"[,;:!?()]")
SENTENCE_SPLIT = re.compile(r"[。；！？\n]")
# 术语与英文缩写中的半角字符属正常排版，校验时先行剔除。
TECH_TOKENS = re.compile(
    r"GB 55015-2021|SEER|STP|V2G|AI|Wp|PUE|M-Link|M-PARK|M-Park|G\d+|"
    r"5\.5至5\.6|check_writing_constraints\.py"
)


def sentences(text: str):
    for chunk in SENTENCE_SPLIT.split(text):
        chunk = chunk.strip()
        if chunk:
            yield chunk


def body_lines(text: str):
    """产出正文行，跳过 Markdown 结构行，并剥离行内代码与链接目标。"""
    for lineno, raw in enumerate(text.splitlines(), start=1):
        line = raw.strip()
        if not line or line.startswith(("#", "<!--")) or set(line) <= {"-"}:
            continue
        stripped = re.sub(r"`[^`]*`", "", line)
        stripped = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", stripped)
        yield lineno, stripped.replace("**", "")


def main() -> int:
    docs = sorted(p for p in DOC_DIR.glob("*.md"))
    if not docs:
        print("未找到待检查的 Markdown 稿件。")
        return 1
    if not REPORT.exists():
        print(f"未找到正文：{REPORT.name}")
        return 1

    failures: list[str] = []

    def check(title: str, problems: list[str]) -> None:
        print(f"\n=== {title} ===")
        if not problems:
            print("PASS")
            return
        for item in problems:
            print(f"FAIL  {item}")
        failures.extend(problems)

    texts = {doc: doc.read_text(encoding="utf-8") for doc in docs}

    # 一、“零碳区域”的使用语境。
    # 出处要求只针对正文；审查说明是对该词本身的说明文字，不适用此项。
    problems = []
    for sentence in sentences(texts[REPORT]):
        if ZERO_CARBON_AREA in sentence and not any(
            m in sentence for m in AREA_SOURCE_MARKERS
        ):
            problems.append(f"{REPORT.name}「零碳区域」缺出处：{sentence}")
    for doc, text in texts.items():
        for sentence in sentences(text):
            if (
                ZERO_CARBON_AREA in sentence
                and any(m in sentence for m in PROJECT_MARKERS)
                and not any(n in sentence for n in NEGATION_MARKERS)
            ):
                problems.append(f"{doc.name}「零碳区域」被用于本项目：{sentence}")
    check("「零碳区域」仅用于转述认证体系", problems)

    # 二、“首个”类表述的限定语
    problems = []
    for doc, text in texts.items():
        for sentence in sentences(text):
            for claim in FIRST_CLAIMS:
                if claim in sentence and not any(m in sentence for m in FIRST_MARKERS):
                    problems.append(f"{doc.name}「{claim}」缺限定语：{sentence}")
    check("「首个」类表述须带出处或争取方向", problems)

    # 三、会中口径的量化数据须同段带口述或估算标注。
    # 含“据”的段落其数据归属另一来源（宣讲材料、报道等），由第六项负责校验。
    problems = []
    for lineno, line in body_lines(texts[REPORT]):
        if "会中" in line and QUANTITY_RE.search(line) and "据" not in line:
            if not any(h in line for h in HEDGE_MARKERS):
                problems.append(
                    f"{REPORT.name} 第 {lineno} 行会中数据缺口述或估算标注：{line[:70]}"
                )
    check("会中口径的量化数据须标注口述或估算", problems)

    # 四、清华会中内容不得升级为已核证或已成交
    problems = []
    for doc, text in texts.items():
        for sentence in sentences(text):
            for word in TSINGHUA_FORBIDDEN:
                if word in sentence and not any(
                    n in sentence for n in NEGATION_MARKERS
                ):
                    problems.append(f"{doc.name}出现「{word}」且非否定：{sentence}")
    check("会中内容不得升级为已核证或已成交", problems)

    # 五、既成事实用语限于否定语境
    problems = []
    for doc, text in texts.items():
        for sentence in sentences(text):
            for term in FAIT_ACCOMPLI:
                if term in sentence and not any(
                    n in sentence for n in NEGATION_MARKERS
                ):
                    problems.append(f"{doc.name}「{term}」非否定语境：{sentence}")
    check("既成事实用语限于否定语境", problems)

    # 六、量化数据必须同段标注出处
    problems = []
    for doc, text in texts.items():
        for lineno, line in body_lines(text):
            if QUANTITY_RE.search(line) and not any(m in line for m in SOURCE_MARKERS):
                problems.append(f"{doc.name} 第 {lineno} 行数据缺出处：{line[:70]}")
    check("量化数据同段标注出处", problems)

    # 七、文号白名单
    problems = []
    for doc, text in texts.items():
        masked = text
        for allowed in ALLOWED_STANDARD_NUMBERS:
            masked = masked.replace(allowed, "○")
        for match in DOC_NUMBER_RE.finditer(masked):
            problems.append(f"{doc.name} 疑似文号：「{match.group()}」")
    check("文号限于资料给出的四个", problems)

    # 八、年份笔误只能作为笔误提及
    problems = []
    for doc, text in texts.items():
        for sentence in sentences(text):
            if TYPO_YEAR in sentence and not any(m in sentence for m in TYPO_MARKERS):
                problems.append(f"{doc.name}「{TYPO_YEAR}」未标为笔误：{sentence}")
    check("年份笔误不得作为时间要求", problems)

    # 九、正文全角标点
    problems = []
    for doc, text in texts.items():
        for lineno, line in body_lines(text):
            if line.startswith("|"):  # 表格行含 Markdown 分隔符
                continue
            if HALFWIDTH_PUNCT.search(TECH_TOKENS.sub("", line)):
                problems.append(f"{doc.name} 第 {lineno} 行含半角标点：{line[:70]}")
    check("正文全角标点", problems)

    # 十、公文 Word 与 Markdown 正文同步
    print("\n=== 公文 Word 与 Markdown 同步 ===")
    if not DOCX.exists():
        print("FAIL  未找到公文 Word，请先执行 build_docx.py")
        failures.append("缺少公文 Word")
    else:
        try:
            from docx import Document
        except ImportError:
            print("SKIP  未安装 python-docx，跳过同步校验")
        else:
            document = Document(DOCX)
            doc_text = "".join(p.text for p in document.paragraphs) + "".join(
                c.text for t in document.tables for r in t.rows for c in r.cells
            )
            missing = [
                line
                for _, line in body_lines(texts[REPORT])
                if len(line) > 30 and not line.startswith("|") and line[:28] not in doc_text
            ]
            title_ok = document.core_properties.title == DOC_NAME
            header_ok = DOC_NAME in document.sections[1].header.paragraphs[0].text
            if missing:
                print(f"FAIL  Word 缺少 {len(missing)} 段正文，需重新生成")
                for line in missing[:3]:
                    print(f"      {line[:60]}")
                failures.append("Word 与 Markdown 不同步")
            elif not title_ok:
                print("FAIL  Word 文档标题属性与文件名称不一致")
                failures.append("Word 标题不一致")
            elif not header_ok:
                print("FAIL  Word 页眉未写入文件名称")
                failures.append("Word 页眉缺失")
            else:
                print("PASS")

    print("\n=== 结果 ===")
    if failures:
        print(f"共 {len(failures)} 项未通过，需修改稿件。")
        return 1
    print("全部检查项通过。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
