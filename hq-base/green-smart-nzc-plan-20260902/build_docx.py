#!/usr/bin/env python3
"""把《南网总部基地绿色智慧近零碳园区规划方案_2026-09-02.md》转成公文体例的 Word 文稿。

用法：
    python3 build_docx.py            # 在本目录生成同名 .docx

只依赖 python-docx（pip install python-docx）。Markdown 只使用有限子集：
# / ## / ### / #### 标题、段落、- 无序列表、1. 有序列表、| 表格 |、--- 分隔线、
**加粗**、`代码`、```mermaid 代码块。

版式：A4；正文仿宋（仿宋_GB2312）三号；一级标题黑体小二居中、二级标题黑体三号、
三级标题仿宋三号加粗；封面标题方正小标宋；表格宋体小五；页脚页码；封面后附目录域。

mermaid 代码块的处理：若 figures/fig-<序号>.png 存在（序号按出现顺序从 1 起），
插入该图并加图题；否则跳过代码块（正文已有文字示意）。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SOURCE = HERE / "南网总部基地绿色智慧近零碳园区规划方案_2026-09-02.md"
OUTPUT = HERE / "南网总部基地绿色智慧近零碳园区规划方案_2026-09-02.docx"
FIGURES = HERE / "figures"

BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
TITLE_FONT = "方正小标宋简体"
TABLE_FONT = "宋体"
DARK = RGBColor(0x1F, 0x1F, 0x1F)

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
FIGURE_TITLES = {
    1: "图 4-1 园区空间结构示意（拓扑关系，不代表实际方位与尺度）",
    2: "图 4-2 统一平台系统架构示意",
    3: "图 7-1 分期建设时序示意",
}


def set_font(run, name: str, size: float, bold: bool | None = None, color: RGBColor = DARK):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = color


def add_inline(paragraph, text: str, size: float, font: str = BODY_FONT, bold_all: bool = False):
    """按 **加粗** 与 `代码` 切分写入 run。"""
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_font(run, font, size, bold=True)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_font(run, TABLE_FONT, size, bold=bold_all)
        else:
            run = paragraph.add_run(piece)
            set_font(run, font, size, bold=bold_all)


def para_format(p, *, first_line_indent=True, align=None, space_after=6, line_spacing=1.5):
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    if first_line_indent:
        fmt.first_line_indent = Pt(24)
    if align is not None:
        p.alignment = align


def set_cell_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tbl_pr.append(borders)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_field(run, instr: str, placeholder: str | None = None):
    """写入域代码；placeholder 为域的当前结果文本，Word 更新域后会被替换。"""
    steps = [("begin", None), (None, instr)]
    if placeholder is not None:
        steps += [("separate", None), ("text", placeholder)]
    steps.append(("end", None))
    for tag, text in steps:
        if tag == "text":
            el = OxmlElement("w:t")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        elif tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(16)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    add_field(run, "PAGE")
    set_font(run, TABLE_FONT, 10.5)
    return doc


def add_toc(doc: Document):
    p = doc.add_paragraph()
    para_format(p, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    p.paragraph_format.page_break_before = True
    run = p.add_run("目　录")
    set_font(run, HEADING_FONT, 16)
    p = doc.add_paragraph()
    para_format(p, first_line_indent=False, space_after=6, line_spacing=1.3)
    run = p.add_run()
    add_field(run, 'TOC \\o "1-2" \\h \\z \\u', placeholder="（目录域：在 Word 中右键“更新域”生成）")
    set_font(run, TABLE_FONT, 12)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def column_widths(rows: list[list[str]], ncols: int, total_cm: float = 15.6) -> list[float]:
    """按各列内容长度分配列宽，短列不占用过多版面。"""
    min_cm = 1.4
    weights = []
    for j in range(ncols):
        lens = [len(INLINE_RE.sub(lambda m: m.group(0).strip("*`"), r[j])) for r in rows if j < len(r)]
        avg = sum(lens) / max(len(lens), 1)
        weights.append(min(max(avg, 3.0), 36.0))
    total = sum(weights)
    spare = total_cm - min_cm * ncols
    return [min_cm + spare * w / total for w in weights]


def add_table(doc: Document, rows: list[list[str]]):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cell_borders(table)
    widths = column_widths(rows, ncols)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for j in range(ncols):
        table.columns[j].width = Cm(widths[j])
    font_size = 9 if ncols <= 5 else 8
    for i, row in enumerate(rows):
        tr_pr = table.rows[i]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if i == 0:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.width = Cm(widths[j])
            cell.text = ""
            p = cell.paragraphs[0]
            para_format(p, first_line_indent=False, space_after=0, line_spacing=1.15)
            if i == 0:
                p.paragraph_format.keep_with_next = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, text, font_size, font=TABLE_FONT, bold_all=(i == 0))
            if i == 0:
                shade_cell(cell, "E7EEF3")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(doc: Document, level: int, text: str):
    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        para_format(p, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.page_break_before = True
        run = p.add_run(text)
        set_font(run, HEADING_FONT, 18, bold=False)
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
        para_format(p, first_line_indent=False, space_after=6)
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(text)
        set_font(run, HEADING_FONT, 16, bold=False)
    else:
        p = doc.add_paragraph(style="Heading 3")
        para_format(p, first_line_indent=False, space_after=4)
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(text)
        set_font(run, BODY_FONT, 16, bold=True)
    p.paragraph_format.keep_with_next = True


def add_figure(doc: Document, index: int):
    png = FIGURES / f"fig-{index}.png"
    if not png.exists():
        return
    p = doc.add_paragraph()
    para_format(p, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p.add_run().add_picture(str(png), width=Cm(15.0))
    cap = doc.add_paragraph()
    para_format(cap, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    run = cap.add_run(FIGURE_TITLES.get(index, f"图 {index}"))
    set_font(run, TABLE_FONT, 10.5)


def build(source: Path = SOURCE, output: Path = OUTPUT) -> Path:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = setup_document()

    i = 0
    cover_done = False
    figure_index = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped.startswith("<!--"):
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            figure_index += 1
            add_figure(doc, figure_index)
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table(block))
            continue

        if stripped.startswith("# ") and not cover_done:
            # 封面：主标题 + 副题 + 说明行，直到第一个 ## 为止；随后插目录。
            p = doc.add_paragraph()
            para_format(p, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            p.paragraph_format.space_before = Pt(110)
            run = p.add_run(stripped[2:].strip())
            set_font(run, TITLE_FONT, 26, bold=True)
            i += 1
            while i < len(lines) and not lines[i].startswith("## "):
                s = lines[i].strip()
                if s and s != "---":
                    p = doc.add_paragraph()
                    is_subtitle = s.startswith("**（")
                    is_note = s.startswith("稿件性质")
                    para_format(
                        p,
                        first_line_indent=is_note,
                        align=WD_ALIGN_PARAGRAPH.CENTER if (is_subtitle or "：" not in s) else WD_ALIGN_PARAGRAPH.LEFT,
                        space_after=10,
                        line_spacing=1.5 if not is_note else 1.3,
                    )
                    if is_note:
                        p.paragraph_format.space_before = Pt(36)
                    add_inline(p, s, 14 if is_subtitle else (12 if is_note else 14))
                i += 1
            cover_done = True
            add_toc(doc)
            continue

        if stripped.startswith("## "):
            add_heading(doc, 1, stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, 2, stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, 3, stripped[5:].strip())
            i += 1
            continue

        m_ol = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_ol:
            p = doc.add_paragraph()
            para_format(p, first_line_indent=False, space_after=4)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-24)
            add_inline(p, f"{m_ol.group(1)}．{m_ol.group(2)}", 16)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph()
            para_format(p, first_line_indent=False, space_after=4)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-12)
            add_inline(p, "▪ " + stripped[2:], 16)
            i += 1
            continue

        # 普通段落（合并连续行）
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", "- ", "---", "<!--", "```")) or re.match(r"^\d+\.\s", nxt):
                break
            buf.append(nxt)
            i += 1
        text = "".join(buf)
        p = doc.add_paragraph()
        is_meta = text.startswith("**") and text.endswith("**") and text.count("**") == 2
        para_format(p, first_line_indent=not is_meta, space_after=6)
        add_inline(p, text, 16)

    doc.save(output)
    return output


if __name__ == "__main__":
    out = build()
    print(f"written: {out}")
