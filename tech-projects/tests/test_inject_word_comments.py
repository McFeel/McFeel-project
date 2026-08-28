from __future__ import annotations

import importlib.util
import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

ROOT = Path(__file__).resolve().parents[2]
SCRIPT = ROOT / "tech-projects/tools/inject_word_comments.py"
SPEC = importlib.util.spec_from_file_location("inject_word_comments", SCRIPT)
assert SPEC and SPEC.loader
injector = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(injector)


class InjectWordCommentsTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.config = json.loads(
            (ROOT / "tech-projects/tools/comments.json").read_text(encoding="utf-8")
        )

    def _fixture(self, path: Path, comments: list[dict[str, object]]) -> None:
        document = Document()
        document.add_paragraph("夹具开头")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格文字保持不变"
        for comment in comments:
            anchor = str(comment["anchors"][0])
            context = " ".join(str(item) for item in comment.get("paragraph_contains", []))
            paragraph = document.add_paragraph()
            paragraph.add_run("前缀")
            paragraph.add_run(anchor)
            paragraph.add_run(f"后缀 {context}")
        document.save(path)

    def test_config_has_required_counts_and_prefixes(self) -> None:
        injector.validate_config(self.config)
        self.assertEqual(34, len(self.config["RD19"]["comments"]))
        self.assertEqual(13, len(self.config["RD10"]["comments"]))

    def test_injects_native_comments_without_changing_text(self) -> None:
        for project in ("RD19", "RD10"):
            with self.subTest(project=project), tempfile.TemporaryDirectory() as temp:
                temp_path = Path(temp)
                source = temp_path / "source.docx"
                output = temp_path / "output.docx"
                comments = self.config[project]["comments"]
                self._fixture(source, comments)

                with zipfile.ZipFile(source) as archive:
                    unchanged_before = archive.read("word/styles.xml")

                matched = injector.inject(source, output, comments)
                report = injector.verify(source, output, comments)

                self.assertEqual(len(comments), len(matched))
                self.assertTrue(report["body_text_unchanged"])
                self.assertEqual(len(comments), report["comment_count"])

                with zipfile.ZipFile(output) as archive:
                    self.assertEqual(unchanged_before, archive.read("word/styles.xml"))
                    comments_root = etree.fromstring(archive.read("word/comments.xml"))
                nodes = comments_root.xpath("./w:comment", namespaces=injector.NS)
                self.assertEqual(
                    {injector.AUTHOR}, {node.get(injector.W + "author") for node in nodes}
                )
                self.assertEqual(
                    {injector.INITIALS},
                    {node.get(injector.W + "initials") for node in nodes},
                )

    def test_rejects_a_precommented_source(self) -> None:
        comments = self.config["RD10"]["comments"]
        with tempfile.TemporaryDirectory() as temp:
            temp_path = Path(temp)
            source = temp_path / "source.docx"
            first_output = temp_path / "first.docx"
            second_output = temp_path / "second.docx"
            self._fixture(source, comments)
            injector.inject(source, first_output, comments)

            with self.assertRaisesRegex(injector.InjectionError, "必须从未批注原稿开始"):
                injector.inject(first_output, second_output, comments)


if __name__ == "__main__":
    unittest.main()
