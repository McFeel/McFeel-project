#!/usr/bin/env python3
"""Inject the approved RD19/RD10 review list as native Word comments.

The source DOCX is treated as immutable input.  Existing text nodes are never
edited or split: comment range markers are inserted immediately around the
existing run(s) that contain each anchor.  All package parts except
``word/document.xml``, its relationships, ``[Content_Types].xml``, and the new
``word/comments.xml`` are copied byte-for-byte.

Run from the repository root:

    python3 tech-projects/tools/inject_word_comments.py
    python3 tech-projects/tools/inject_word_comments.py --project RD19
    python3 tech-projects/tools/inject_word_comments.py --check-config
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import re
import sys
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from lxml import etree

REPO_ROOT = Path(__file__).resolve().parents[2]
CONFIG_PATH = Path(__file__).with_name("comments.json")

AUTHOR = "科技项目专责"
INITIALS = "科"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
XML_NS = "http://www.w3.org/XML/1998/namespace"

NS = {"w": W_NS}
W = f"{{{W_NS}}}"
CT = "[Content_Types].xml"
DOCUMENT = "word/document.xml"
DOCUMENT_RELS = "word/_rels/document.xml.rels"
COMMENTS = "word/comments.xml"
COMMENTS_CT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
COMMENTS_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
ALLOWED_CHANGED_PARTS = {CT, DOCUMENT, DOCUMENT_RELS, COMMENTS}


class InjectionError(RuntimeError):
    """Raised when source contents cannot be mapped safely to the review list."""


def _load_config() -> dict[str, Any]:
    with CONFIG_PATH.open(encoding="utf-8") as handle:
        return json.load(handle)


def validate_config(config: dict[str, Any]) -> None:
    expected_projects = {"RD19", "RD10"}
    if set(config) != expected_projects:
        raise InjectionError(
            f"comments.json 项目必须恰为 {sorted(expected_projects)}，实际为 {sorted(config)}"
        )

    for project, data in config.items():
        comments = data["comments"]
        if len(comments) != data["expected_count"]:
            raise InjectionError(
                f"{project}: expected_count={data['expected_count']}，"
                f"但清单有 {len(comments)} 条"
            )
        ids = [item["id"] for item in comments]
        if len(ids) != len(set(ids)):
            raise InjectionError(f"{project}: 批注编号重复")
        for item in comments:
            if not item.get("anchors"):
                raise InjectionError(f"{project}/{item['id']}: 缺少锚点")
            if item["id"] != "OVERVIEW" and not re.match(
                r"^\[(高|中|低)\]\s", item["text"]
            ):
                raise InjectionError(
                    f"{project}/{item['id']}: 批注须以 [高]/[中]/[低] 开头"
                )


def _run_text(run: etree._Element) -> str:
    pieces: list[str] = []
    for node in run.iter():
        if node.tag == W + "t":
            pieces.append(node.text or "")
        elif node.tag == W + "tab":
            pieces.append("\t")
        elif node.tag in {W + "br", W + "cr"}:
            pieces.append("\n")
        elif node.tag == W + "noBreakHyphen":
            pieces.append("\N{NON-BREAKING HYPHEN}")
        elif node.tag == W + "softHyphen":
            pieces.append("\N{SOFT HYPHEN}")
    return "".join(pieces)


def _paragraph_runs(paragraph: etree._Element) -> list[etree._Element]:
    return paragraph.xpath(".//w:r", namespaces=NS)


def _paragraph_text(paragraph: etree._Element) -> str:
    return "".join(_run_text(run) for run in _paragraph_runs(paragraph))


def _find_anchor(
    root: etree._Element, spec: dict[str, Any]
) -> tuple[etree._Element, int, int, str]:
    occurrence = int(spec.get("occurrence", 1))
    if occurrence < 1:
        raise InjectionError(f"{spec['id']}: occurrence 必须从 1 开始")
    required_context = spec.get("paragraph_contains", [])

    for anchor in spec["anchors"]:
        hits: list[tuple[etree._Element, int]] = []
        for paragraph in root.xpath(".//w:body//w:p", namespaces=NS):
            text = _paragraph_text(paragraph)
            if any(fragment not in text for fragment in required_context):
                continue
            start = 0
            while True:
                index = text.find(anchor, start)
                if index < 0:
                    break
                hits.append((paragraph, index))
                start = index + max(1, len(anchor))
        if len(hits) >= occurrence:
            paragraph, start = hits[occurrence - 1]
            return paragraph, start, start + len(anchor), anchor

    context_note = f"，段落还须含 {required_context!r}" if required_context else ""
    raise InjectionError(
        f"{spec['id']}: 找不到第 {occurrence} 个锚点 {spec['anchors']!r}{context_note}"
    )


def _runs_for_range(
    paragraph: etree._Element, start: int, end: int
) -> tuple[etree._Element, etree._Element]:
    first: etree._Element | None = None
    last: etree._Element | None = None
    position = 0
    for run in _paragraph_runs(paragraph):
        run_length = len(_run_text(run))
        run_end = position + run_length
        if run_length and run_end > start and position < end:
            first = first or run
            last = run
        position = run_end
    if first is None or last is None:
        raise InjectionError("锚点未覆盖任何文字 run")
    if first.getparent() is not last.getparent():
        raise InjectionError("锚点跨越不同 XML 容器，无法安全插入批注范围")
    return first, last


def _insert_comment_range(
    paragraph: etree._Element, start: int, end: int, comment_id: int
) -> None:
    first, last = _runs_for_range(paragraph, start, end)

    range_start = etree.Element(W + "commentRangeStart")
    range_start.set(W + "id", str(comment_id))
    first.addprevious(range_start)

    range_end = etree.Element(W + "commentRangeEnd")
    range_end.set(W + "id", str(comment_id))
    last.addnext(range_end)

    reference_run = etree.Element(W + "r")
    run_properties = etree.SubElement(reference_run, W + "rPr")
    run_style = etree.SubElement(run_properties, W + "rStyle")
    run_style.set(W + "val", "CommentReference")
    reference = etree.SubElement(reference_run, W + "commentReference")
    reference.set(W + "id", str(comment_id))
    range_end.addnext(reference_run)


def _comments_xml(comments: list[dict[str, Any]]) -> bytes:
    root = etree.Element(W + "comments", nsmap={"w": W_NS})
    timestamp = (
        datetime.now(timezone.utc)
        .replace(microsecond=0)
        .isoformat()
        .replace("+00:00", "Z")
    )
    for comment_id, spec in enumerate(comments):
        comment = etree.SubElement(root, W + "comment")
        comment.set(W + "id", str(comment_id))
        comment.set(W + "author", AUTHOR)
        comment.set(W + "initials", INITIALS)
        comment.set(W + "date", timestamp)

        paragraph = etree.SubElement(comment, W + "p")
        properties = etree.SubElement(paragraph, W + "pPr")
        style = etree.SubElement(properties, W + "pStyle")
        style.set(W + "val", "CommentText")

        annotation_run = etree.SubElement(paragraph, W + "r")
        annotation_properties = etree.SubElement(annotation_run, W + "rPr")
        annotation_style = etree.SubElement(annotation_properties, W + "rStyle")
        annotation_style.set(W + "val", "CommentReference")
        etree.SubElement(annotation_run, W + "annotationRef")

        text_run = etree.SubElement(paragraph, W + "r")
        text = etree.SubElement(text_run, W + "t")
        text.set(f"{{{XML_NS}}}space", "preserve")
        text.text = spec["text"]

    return etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )


def _register_content_type(raw: bytes) -> bytes:
    root = etree.fromstring(raw)
    matches = root.xpath(
        "./ct:Override[@PartName='/word/comments.xml']",
        namespaces={"ct": CT_NS},
    )
    if not matches:
        override = etree.SubElement(root, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", COMMENTS_CT)
    return etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )


def _register_relationship(raw: bytes | None) -> bytes:
    if raw is None:
        root = etree.Element(f"{{{REL_NS}}}Relationships")
    else:
        root = etree.fromstring(raw)
    existing = root.xpath(
        "./pr:Relationship[@Type=$relationship_type]",
        namespaces={"pr": REL_NS},
        relationship_type=COMMENTS_REL,
    )
    if not existing:
        used = {
            relation.get("Id")
            for relation in root.findall(f"{{{REL_NS}}}Relationship")
        }
        number = 1
        while f"rId{number}" in used:
            number += 1
        relation = etree.SubElement(root, f"{{{REL_NS}}}Relationship")
        relation.set("Id", f"rId{number}")
        relation.set("Type", COMMENTS_REL)
        relation.set("Target", "comments.xml")
    return etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )


def _copy_zip_info(info: zipfile.ZipInfo) -> zipfile.ZipInfo:
    copied = copy.copy(info)
    copied.CRC = 0
    copied.file_size = 0
    copied.compress_size = 0
    return copied


def inject(source: Path, output: Path, comments: list[dict[str, Any]]) -> list[str]:
    with zipfile.ZipFile(source, "r") as archive:
        infos = archive.infolist()
        if len({info.filename for info in infos}) != len(infos):
            raise InjectionError(f"{source}: DOCX 含重复 ZIP 部件名")
        parts = {info.filename: archive.read(info.filename) for info in infos}

    for required in (CT, DOCUMENT):
        if required not in parts:
            raise InjectionError(f"{source}: 缺少 {required}，不是有效 DOCX")
    if COMMENTS in parts:
        raise InjectionError(f"{source}: 已有 {COMMENTS}；必须从未批注原稿开始")

    document_root = etree.fromstring(parts[DOCUMENT])
    if document_root.xpath(
        ".//w:commentRangeStart | .//w:commentRangeEnd | .//w:commentReference",
        namespaces=NS,
    ):
        raise InjectionError(f"{source}: 正文已有批注标记；必须从未批注原稿开始")

    matched: list[str] = []
    for comment_id, spec in enumerate(comments):
        paragraph, start, end, anchor = _find_anchor(document_root, spec)
        _insert_comment_range(paragraph, start, end, comment_id)
        matched.append(anchor)

    parts[DOCUMENT] = etree.tostring(
        document_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    parts[COMMENTS] = _comments_xml(comments)
    parts[CT] = _register_content_type(parts[CT])
    parts[DOCUMENT_RELS] = _register_relationship(parts.get(DOCUMENT_RELS))

    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        prefix=output.name + ".", suffix=".tmp", dir=output.parent, delete=False
    ) as temp:
        temp_path = Path(temp.name)
    try:
        info_by_name = {info.filename: info for info in infos}
        ordered_names = [info.filename for info in infos]
        for added_name in (DOCUMENT_RELS, COMMENTS):
            if added_name not in ordered_names:
                ordered_names.append(added_name)

        with zipfile.ZipFile(temp_path, "w") as archive:
            for name in ordered_names:
                if name in info_by_name:
                    archive.writestr(_copy_zip_info(info_by_name[name]), parts[name])
                else:
                    archive.writestr(name, parts[name], compress_type=zipfile.ZIP_DEFLATED)
        temp_path.replace(output)
    finally:
        temp_path.unlink(missing_ok=True)
    return matched


def _document_text_signature(path: Path) -> dict[str, Any]:
    document = Document(path)

    def table_signature(table: Any) -> list[Any]:
        rows: list[Any] = []
        for row in table.rows:
            cells: list[Any] = []
            for cell in row.cells:
                cells.append(
                    {
                        "paragraphs": [paragraph.text for paragraph in cell.paragraphs],
                        "tables": [table_signature(nested) for nested in cell.tables],
                    }
                )
            rows.append(cells)
        return rows

    return {
        "paragraphs": [paragraph.text for paragraph in document.paragraphs],
        "tables": [table_signature(table) for table in document.tables],
    }


def _sha256(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def verify(
    source: Path, output: Path, comments: list[dict[str, Any]]
) -> dict[str, Any]:
    source_signature = _document_text_signature(source)
    output_signature = _document_text_signature(output)
    if output_signature != source_signature:
        raise InjectionError(f"{output}: python-docx 段落/表格文字与原稿不一致")

    with zipfile.ZipFile(source) as source_zip, zipfile.ZipFile(output) as output_zip:
        source_parts = {name: source_zip.read(name) for name in source_zip.namelist()}
        output_parts = {name: output_zip.read(name) for name in output_zip.namelist()}

    for name, raw in source_parts.items():
        if name not in ALLOWED_CHANGED_PARTS:
            if name not in output_parts or output_parts[name] != raw:
                raise InjectionError(f"{output}: 非批注部件 {name} 未保持原始字节")

    comments_root = etree.fromstring(output_parts[COMMENTS])
    comment_nodes = comments_root.xpath("./w:comment", namespaces=NS)
    starts = etree.fromstring(output_parts[DOCUMENT]).xpath(
        ".//w:commentRangeStart", namespaces=NS
    )
    ends = etree.fromstring(output_parts[DOCUMENT]).xpath(
        ".//w:commentRangeEnd", namespaces=NS
    )
    references = etree.fromstring(output_parts[DOCUMENT]).xpath(
        ".//w:commentReference", namespaces=NS
    )
    expected_count = len(comments)
    counts = {
        "comments.xml": len(comment_nodes),
        "commentRangeStart": len(starts),
        "commentRangeEnd": len(ends),
        "commentReference": len(references),
    }
    if any(count != expected_count for count in counts.values()):
        raise InjectionError(f"{output}: 批注部件计数错误：{counts}")

    expected_ids = {str(number) for number in range(expected_count)}
    for label, nodes in (
        ("comment", comment_nodes),
        ("start", starts),
        ("end", ends),
        ("reference", references),
    ):
        ids = {node.get(W + "id") for node in nodes}
        if ids != expected_ids:
            raise InjectionError(f"{output}: {label} id 不闭合：{sorted(ids)}")
    if any(
        node.get(W + "author") != AUTHOR or node.get(W + "initials") != INITIALS
        for node in comment_nodes
    ):
        raise InjectionError(f"{output}: 批注作者或 initials 不正确")

    return {
        "source": str(source.relative_to(REPO_ROOT)),
        "output": str(output.relative_to(REPO_ROOT)),
        "comment_count": expected_count,
        "body_text_unchanged": True,
        "source_sha256": _sha256(source.read_bytes()),
        "output_sha256": _sha256(output.read_bytes()),
    }


def _selected_projects(
    config: dict[str, Any], requested: Iterable[str]
) -> list[tuple[str, dict[str, Any]]]:
    names = list(requested)
    return [(name, config[name]) for name in (names or ["RD19", "RD10"])]


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--project", action="append", choices=("RD19", "RD10"), default=[]
    )
    parser.add_argument(
        "--check-config",
        action="store_true",
        help="只校验批注清单，不要求原稿存在",
    )
    args = parser.parse_args(argv)

    try:
        config = _load_config()
        validate_config(config)
        if args.check_config:
            for name, data in _selected_projects(config, args.project):
                print(f"{name}: {len(data['comments'])} 条批注，配置有效")
            return 0

        missing: list[Path] = []
        selected = _selected_projects(config, args.project)
        for _, data in selected:
            source = REPO_ROOT / data["source"]
            if not source.is_file():
                missing.append(source)
        if missing:
            print("缺少原始用户 Word 文件；未生成任何批注稿：", file=sys.stderr)
            for path in missing:
                print(f"  - {path.relative_to(REPO_ROOT)}", file=sys.stderr)
            print("请把原始二进制文件放到以上固定路径后重跑本命令。", file=sys.stderr)
            return 1

        for name, data in selected:
            source = REPO_ROOT / data["source"]
            output = REPO_ROOT / data["output"]
            matched = inject(source, output, data["comments"])
            report = verify(source, output, data["comments"])
            print(
                f"{name}: 已生成 {report['output']}；"
                f"{report['comment_count']} 条；正文文字不变"
            )
            for spec, anchor in zip(data["comments"], matched, strict=True):
                print(f"  {spec['id']:>8} -> {anchor}")
        return 0
    except (InjectionError, OSError, ValueError, zipfile.BadZipFile) as error:
        print(f"错误：{error}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
